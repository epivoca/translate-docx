from docx import Document
import mtranslate


def add_formats_to_run(orig_run, new_run):
    # set the style of the new run to match the original run
    # TODO check if there any other styles in the documentation
    new_run.bold = orig_run.bold
    new_run.italic = orig_run.italic
    new_run.underline = orig_run.underline
    new_run.font.strike = orig_run.font.strike
    new_run.font.subscript = orig_run.font.subscript
    new_run.font.superscript = orig_run.font.superscript
    new_run.font.color.rgb = orig_run.font.color.rgb


def translate_paragraphs(doc_paragraphs, translated_doc, style_map):
    for orig_paragraph in doc_paragraphs:
        # add a new paragraph to the translated document
        new_paragraph = translated_doc.add_paragraph('')

        # set the style of the new paragraph to match the original paragraph
        new_paragraph.style = orig_paragraph.style

        # iterate through each run in the original paragraph
        for orig_run in orig_paragraph.runs:
            # add a new run to the translated paragraph
            new_run = new_paragraph.add_run('')

            add_formats_to_run(orig_run=orig_run, new_run=new_run)

            # TODO remove translator to the class
            run_translation = mtranslate.translate(
                from_language='en',
                to_language='ru',
                to_translate=orig_run.text
            )

            # set the text of the new run to the translated text
            new_run.text = run_translation

        # add the style of the original paragraph to the style map
        style_map[orig_paragraph.style.name] = new_paragraph.style.name


def translate_tables(doc, translated_doc, style_map):
    for i, orig_table in enumerate(doc.tables):
        # TODO add preserving of table borders
        # TODO think how to save table position in the document
        new_table = translated_doc.add_table(rows=len(orig_table.rows), cols=len(orig_table.columns))

        # copy the border settings from the original table to the new table
        new_table.style = orig_table.style
        new_table.alignment = orig_table.alignment
        new_table.autofit = orig_table.autofit

        # set the width of each column in the new table to match the original table
        for j, orig_col in enumerate(orig_table.columns):
            new_table.columns[j].width = orig_col.width

        for k, orig_cell in enumerate(orig_table._cells):
            # get the corresponding cell in the new table
            new_cell = new_table.cell(k // len(orig_table.columns), k % len(orig_table.columns))
            translate_paragraphs(doc_paragraphs=doc.paragraphs, translated_doc=new_cell, style_map=style_map)


def translate_docx():
    # TODO remove to func parameters
    doc = Document('input.docx')

    translated_doc = Document()

    # dictionary to map paragraph styles between documents
    style_map = {}

    translate_paragraphs(doc_paragraphs=doc.paragraphs, translated_doc=translated_doc, style_map=style_map)
    translate_tables(doc=doc, translated_doc=translated_doc, style_map=style_map)

    translated_doc.save('translated_document.docx')
